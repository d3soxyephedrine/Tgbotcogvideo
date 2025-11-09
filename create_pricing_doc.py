#!/usr/bin/env python3
"""Create a professional pricing document for the Telegram bot"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_row(table, cells, bold=False, header=False):
    """Add a row to a table"""
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(cell_text)
        if bold or header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        if header:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_pricing_document():
    """Create the pricing document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('UnboundGPT Bot - Pricing & Features Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # ========== TEXT GENERATION ==========
    add_heading(doc, '💬 Text Generation (Chat)', level=1)
    
    # Create table for text generation
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Model', 'Command', 'Credits', 'Details'], header=True)
    add_table_row(table, ['DeepSeek-Chat-V3-0324', 'Default', '1 credit', 'Cost-effective, uncensored AI via OpenRouter'])
    add_table_row(table, ['ChatGPT-4o-latest', 'Toggle via /model', '3 credits', 'Premium reasoning model via OpenRouter'])
    add_table_row(table, ['Writing Mode', '/write <prompt>', '2 credits', 'Professional prose, minimum 300 words (always 2 credits)'])
    
    doc.add_paragraph()
    
    # ========== IMAGE GENERATION ==========
    add_heading(doc, '🎨 Image Generation', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Model', 'Command', 'Credits', 'Provider'], header=True)
    add_table_row(table, ['FLUX.1 Kontext Max', '/imagine <prompt>', '10 credits', 'Novita AI - photorealistic'])
    add_table_row(table, ['Hunyuan-Image-3', '/uncensored <prompt>', '10 credits', 'Novita AI - fully uncensored'])
    add_table_row(table, ['Grok-2-Image', '/grok <prompt>', '8 credits', 'xAI API - stylized'])
    add_table_row(table, ['Qwen-Image', '/edit <prompt>', '8 credits', 'Novita AI - great for text/editing'])
    
    doc.add_paragraph()
    
    # ========== IMAGE EDITING ==========
    add_heading(doc, '✨ Image Editing (requires uploaded photo)', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Model', 'Command', 'Credits', 'Provider'], header=True)
    add_table_row(table, ['FLUX.1 Kontext Max', 'Photo + caption', '15 credits', 'Novita AI - maximum permissiveness'])
    add_table_row(table, ['Qwen-Image', 'Photo + /edit <caption>', '12 credits', 'Novita AI - standard editing'])
    
    doc.add_paragraph()
    
    # ========== VIDEO GENERATION ==========
    add_heading(doc, '🎬 Video Generation (Image-to-Video)', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Model', 'Command', 'Credits', 'Provider'], header=True)
    add_table_row(table, ['WAN 2.5 I2V Preview', 'Photo + /img2video <caption>', '50 credits', 'Novita AI - async processing'])
    
    doc.add_paragraph()
    
    # ========== FREE FEATURES ==========
    add_heading(doc, '🆓 Free Features (0 Credits)', level=1)
    
    free_features = [
        '! memorize <text> - Store memory',
        '! memories - List all memories',
        '! forget <id> - Delete memory',
        '/balance - Check credits',
        '/help - Show help',
        '/model - Switch models',
        '/daily - Claim free daily credits',
        '/clear - Clear history'
    ]
    
    for feature in free_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_paragraph()
    
    # ========== QUICK SUMMARY ==========
    add_heading(doc, '📊 Quick Summary', level=1)
    
    summary_text = [
        ('Cheapest:', 'Text chat (1 credit), Memory commands (FREE)'),
        ('Mid-range:', 'Image generation (8-10 credits), Writing mode (2 credits)'),
        ('Premium:', 'Image editing (12-15 credits), Video generation (50 credits)')
    ]
    
    for label, info in summary_text:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f' {info}')
    
    doc.add_page_break()
    
    # ========== CREDIT PACKAGES ==========
    add_heading(doc, '💰 Credit Packages', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    add_table_row(table, ['Package', 'Credits', 'Price per Credit', 'Total Cost'], header=True)
    add_table_row(table, ['Starter', '200 credits', '2.5¢', '$5.00'])
    add_table_row(table, ['Popular', '400 credits', '2.5¢', '$10.00'])
    add_table_row(table, ['Value', '800 credits', '2.5¢', '$20.00'])
    add_table_row(table, ['Premium', '2,000 credits', '2.5¢', '$50.00'])
    
    doc.add_paragraph()
    
    # ========== FREE CREDITS ==========
    add_heading(doc, '🎁 Free Credits', level=1)
    
    free_credits_info = [
        ('New users:', '100 free credits (for text chat only)'),
        ('Daily credits:', '25 free credits via /daily command'),
        ('Frequency:', 'Claimable once every 24 hours'),
        ('Expiration:', 'Expires after 48 hours'),
        ('Usage priority:', 'Used before purchased credits')
    ]
    
    for label, info in free_credits_info:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f' {info}')
    
    doc.add_paragraph()
    
    # ========== EXAMPLE CALCULATIONS ==========
    add_heading(doc, '📊 Example Calculations', level=1)
    
    doc.add_paragraph().add_run('With the $5 package (200 credits), you could get:').bold = True
    
    examples = [
        '200 text messages (DeepSeek)',
        '66 text messages (GPT-4o)',
        '100 writing requests (/write)',
        '20 images (/imagine or /uncensored)',
        '25 images (/grok or /edit)',
        '13 image edits (FLUX)',
        '16 image edits (Qwen)',
        '4 videos (/img2video)'
    ]
    
    for example in examples:
        doc.add_paragraph(example, style='List Bullet')
    
    doc.add_paragraph()
    note = doc.add_paragraph('Note: You can mix and match any combination of features based on your needs!')
    note.runs[0].italic = True
    
    # Save document
    filename = 'UnboundGPT_Pricing_Guide.docx'
    doc.save(filename)
    print(f"✅ Document created successfully: {filename}")
    return filename

if __name__ == '__main__':
    create_pricing_document()
