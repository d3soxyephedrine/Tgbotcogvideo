import os
import logging
import threading
import time
import requests
import json
import hmac
import hashlib
import uuid
from urllib.parse import urlparse
from flask import Flask, request, jsonify, render_template_string, render_template
from telegram_handler import process_update, send_message
from llm_api import generate_response, OPENROUTER_API_KEY, OPENROUTER_ENDPOINT
from models import db, User, Message, Payment, Transaction, CryptoPayment, Conversation, TelegramPayment
from datetime import datetime
from sqlalchemy import desc
from nowpayments_api import NOWPaymentsAPI
from nowpayments_wrapper import NOWPaymentsWrapper
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import send_file
import io

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Global flag to track database availability
DB_AVAILABLE = False
DB_INIT_ATTEMPTS = 0
MAX_DB_INIT_ATTEMPTS = 3

# Track whether we're using the on-disk SQLite fallback database
USING_SQLITE_FALLBACK = False


def ensure_database_url():
    """Ensure DATABASE_URL is configured, falling back to SQLite if needed."""

    global USING_SQLITE_FALLBACK

    database_url = os.environ.get("DATABASE_URL")
    if database_url:
        return database_url

    # No DATABASE_URL configured – fall back to a local SQLite database so
    # credit-related commands (/daily, /balance, etc.) still function in
    # development environments.
    fallback_path = os.path.join(os.path.dirname(__file__), "local.db")
    os.makedirs(os.path.dirname(fallback_path), exist_ok=True)
    database_url = f"sqlite:///{fallback_path}"
    os.environ["DATABASE_URL"] = database_url
    USING_SQLITE_FALLBACK = True

    logging.getLogger(__name__).warning(
        "DATABASE_URL not set – using fallback SQLite database at %s", fallback_path
    )

    return database_url

DATABASE_URL = ensure_database_url()


def validate_environment():
    """Validate required environment variables at startup"""
    required_vars = ["BOT_TOKEN"]
    optional_vars = ["DATABASE_URL", "NOVITA_API_KEY", "OPENROUTER_API_KEY", "MODEL", "SESSION_SECRET"]
    
    missing_required = [var for var in required_vars if not os.environ.get(var)]
    missing_optional = [var for var in optional_vars if not os.environ.get(var)]
    
    if missing_required:
        logger.error(f"CRITICAL: Missing required environment variables: {', '.join(missing_required)}")
        logger.error("Application may not function correctly without these variables")
    
    if missing_optional:
        logger.warning(f"Optional environment variables not set: {', '.join(missing_optional)}")
        logger.warning("Some features may be disabled")
    
    # Log database configuration status
    if os.environ.get("DATABASE_URL"):
        if USING_SQLITE_FALLBACK:
            logger.info("DATABASE_URL configured via fallback SQLite database")
        else:
            logger.info("DATABASE_URL is configured")
    else:
        logger.warning("DATABASE_URL not configured - database features will be disabled")

    return len(missing_required) == 0

# Validate environment variables
env_valid = validate_environment()

# Initialize Flask application
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "default_secret_key")

# Configure the database with Cloud Run compatible settings
if DATABASE_URL:
    app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL

    parsed_db_url = urlparse(DATABASE_URL)
    is_sqlite = parsed_db_url.scheme.startswith("sqlite")

    if is_sqlite:
        # SQLite does not support the same engine options as Postgres.
        app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
            "connect_args": {"check_same_thread": False}
        }
        logger.info("Configured SQLite engine options (fallback database)")
    else:
        app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
            "pool_recycle": 300,
            "pool_pre_ping": True,
            "pool_size": 10,
            "max_overflow": 20,
            "pool_timeout": 30,
            "connect_args": {
                "connect_timeout": 10,
                "options": "-c statement_timeout=30000"
            }
        }
        logger.info("Configured production database engine options")

    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    logger.info("Database configuration applied")
else:
    logger.warning("Skipping database configuration - DATABASE_URL not set")

def validate_database_connection():
    """Validate database connection without blocking startup"""
    global DB_AVAILABLE, DB_INIT_ATTEMPTS
    
    if not DATABASE_URL:
        logger.warning("Database connection skipped - DATABASE_URL not configured")
        return False
    
    try:
        with app.app_context():
            # Test connection
            db.engine.connect()
            logger.info("Database connection validated successfully")
            return True
    except Exception as e:
        logger.error(f"Database connection validation failed: {str(e)}")
        return False

def init_database():
    """Initialize database tables safely with retry logic"""
    global DB_AVAILABLE, DB_INIT_ATTEMPTS
    
    if not DATABASE_URL:
        logger.info("Database initialization skipped - no DATABASE_URL configured")
        return
    
    max_retries = MAX_DB_INIT_ATTEMPTS
    retry_delay = 2
    
    for attempt in range(1, max_retries + 1):
        DB_INIT_ATTEMPTS = attempt
        try:
            logger.info(f"Database initialization attempt {attempt}/{max_retries}")
            
            # Initialize database connection
            db.init_app(app)
            
            with app.app_context():
                # Validate connection first
                db.engine.connect()
                
                # Create tables
                db.create_all()
                
                # Mark as available
                DB_AVAILABLE = True
                
                # Sync with telegram_handler
                try:
                    from telegram_handler import set_db_available
                    set_db_available(True)
                except ImportError:
                    logger.warning("Could not sync database status with telegram_handler")
                
                logger.info("Database tables created successfully")
                return
                
        except Exception as e:
            logger.error(f"Database initialization attempt {attempt} failed: {str(e)}")
            
            if attempt < max_retries:
                logger.info(f"Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                logger.error("Database initialization failed after all retries")
                logger.warning("App will continue without database - some features will be disabled")

# Initialize database synchronously (blocking) to ensure DB_AVAILABLE is set before handling requests
if DATABASE_URL:
    init_database()
else:
    logger.info("Database initialization skipped - running without database")

# Get environment variables
BOT_TOKEN = os.environ.get("BOT_TOKEN")
if not BOT_TOKEN:
    logger.error("BOT_TOKEN environment variable is not set!")

# Initialize NOWPayments with custom wrapper
NOWPAYMENTS_API_KEY = os.environ.get("NOWPAYMENTS_API_KEY")
if NOWPAYMENTS_API_KEY:
    nowpayments = NOWPaymentsWrapper(NOWPAYMENTS_API_KEY)
    logger.info("NOWPayments API configured (using custom wrapper)")
else:
    nowpayments = None
    logger.warning("NOWPAYMENTS_API_KEY not set - crypto payment features will be disabled")

NOWPAYMENTS_IPN_SECRET = os.environ.get("NOWPAYMENTS_IPN_SECRET")
if NOWPAYMENTS_IPN_SECRET:
    logger.info("NOWPayments IPN secret configured")
else:
    logger.warning("NOWPAYMENTS_IPN_SECRET not set - IPN callbacks will not be verified")

# Keepalive URL (pings local server to keep it awake)
# Use the same port that gunicorn is bound to (from PORT env var or default 5000)
KEEPALIVE_PORT = os.environ.get("PORT", "5000")
KEEPALIVE_URL = f"http://localhost:{KEEPALIVE_PORT}"

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def auto_title_conversation_if_first_message(conversation, user_message, conversation_id):
    """Auto-title conversation if this is the first message"""
    message_count = Message.query.filter_by(conversation_id=conversation_id).count()
    if message_count == 0 and conversation.title == 'New Chat':
        new_title = user_message[:50]
        if len(user_message) > 50:
            new_title += "..."
        conversation.title = new_title
        logger.info(f"Auto-titled conversation {conversation_id}: {new_title}")

def mask_api_key(key):
    """Mask API key for safe logging"""
    if not key or len(key) < 8:
        return "***"
    return f"{key[:3]}...{key[-3:]}"

from functools import wraps

def require_api_key(f):
    """Decorator to require and validate API key authentication"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not DB_AVAILABLE:
            return jsonify({"error": "Service temporarily unavailable"}), 503
            
        auth_header = request.headers.get('Authorization', '')
        if not auth_header.startswith('Bearer '):
            return jsonify({"error": "Missing or invalid authorization header"}), 401
        
        api_key = auth_header[7:]
        user = User.query.filter_by(api_key=api_key).first()
        
        if not user:
            logger.warning(f"Invalid API key attempted: {mask_api_key(api_key)}")
            return jsonify({"error": "Invalid API key"}), 401
        
        # Pass user to the route function as keyword argument
        kwargs['user'] = user
        return f(*args, **kwargs)
    return decorated_function

# ============================================================================
# ROUTES
# ============================================================================

@app.route('/')
def home():
    """Landing page for Ko2 bot"""
    return render_template('index.html')

@app.route('/health')
def health_check():
    """Comprehensive health check endpoint with database latency testing and memory monitoring"""
    import os
    import psutil
    start_time = time.time()
    
    health_status = {
        "status": "healthy",
        "environment_valid": env_valid,
        "database": {
            "configured": DATABASE_URL is not None,
            "available": DB_AVAILABLE,
            "init_attempts": DB_INIT_ATTEMPTS
        },
        "bot_token_configured": BOT_TOKEN is not None,
        "timestamp": time.time()
    }
    
    # Memory monitoring
    try:
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        health_status["memory"] = {
            "rss_mb": round(memory_info.rss / 1024 / 1024, 2),  # Resident Set Size
            "vms_mb": round(memory_info.vms / 1024 / 1024, 2),  # Virtual Memory Size
            "percent": round(process.memory_percent(), 2)
        }
        
        # Warn if memory usage is high (>400MB per worker)
        if memory_info.rss > 400 * 1024 * 1024:
            health_status["memory"]["warning"] = "High memory usage"
            health_status["status"] = "warning"
    except Exception as e:
        logger.warning(f"Memory monitoring failed: {str(e)}")
        health_status["memory"] = {"error": "Unable to get memory stats"}
    
    # Test database latency if available
    if DB_AVAILABLE:
        try:
            db_start = time.time()
            # Simple query to test database responsiveness
            db.session.execute(db.text("SELECT 1"))
            db_latency_ms = (time.time() - db_start) * 1000
            
            health_status["database"]["latency_ms"] = round(db_latency_ms, 2)
            health_status["database"]["status"] = "responsive"
            
            # Get table counts for monitoring
            user_count = User.query.count()
            message_count = Message.query.count()
            conversation_count = Conversation.query.count()
            
            health_status["database"]["stats"] = {
                "users": user_count,
                "messages": message_count,
                "conversations": conversation_count
            }
            
            # Warn if database latency is high
            if db_latency_ms > 100:
                health_status["database"]["warning"] = "High latency"
        except Exception as e:
            logger.error(f"Database health check failed: {str(e)}")
            health_status["database"]["status"] = "error"
            health_status["database"]["error"] = str(e)
            health_status["status"] = "degraded"
    
    # Calculate total response time
    total_time_ms = (time.time() - start_time) * 1000
    health_status["response_time_ms"] = round(total_time_ms, 2)
    
    # Return 503 if critical components are missing
    status_code = 200
    if not env_valid or not BOT_TOKEN:
        health_status["status"] = "degraded"
        status_code = 503
    elif DATABASE_URL and not DB_AVAILABLE:
        health_status["status"] = "degraded"
        health_status["database"]["message"] = "Database configured but not available"
    
    return jsonify(health_status), status_code

@app.route('/chat')
def chat():
    """Web chat interface for LibreChat-like experience"""
    return render_template('chat.html')

@app.route('/prompt-maker')
def prompt_maker():
    """Canmore Syntax Injector - Base64 payload generator"""
    return render_template('prompt-maker.html')

@app.route('/api/balance', methods=['GET'])
@require_api_key
def get_balance(**kwargs):
    user = kwargs['user']
    """Get user's credit balance (authenticated via API key)"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Service temporarily unavailable"
        }), 503
    
    try:
        # Calculate total credits with defensive handling for None values
        daily_credits = user.daily_credits if user.daily_credits is not None else 0
        purchased_credits = user.credits if user.credits is not None else 0
        total_credits = daily_credits + purchased_credits
        
        return jsonify({
            "daily_credits": daily_credits,
            "purchased_credits": purchased_credits,
            "total_credits": total_credits
        })
    except Exception as e:
        logger.error(f"Error fetching balance: {str(e)}")
        return jsonify({
            "error": "Internal server error"
        }), 500

@app.route('/api/messages', methods=['GET'])
@require_api_key
def get_messages(**kwargs):
    user = kwargs['user']
    """Get user's web chat message history (authenticated via API key)
    
    Query params:
        conversation_id (optional): Filter messages by conversation ID
    """
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Service temporarily unavailable"
        }), 503
    
    conversation_id = request.args.get('conversation_id', type=int)
    
    try:
        # Get web messages for this user, optionally filtered by conversation
        from sqlalchemy import desc
        query = Message.query.filter_by(
            user_id=user.id,
            platform='web'
        )
        
        # Filter by conversation if specified
        if conversation_id is not None:
            # When loading a specific conversation, load ALL messages (no limit)
            query = query.filter_by(conversation_id=conversation_id)
            messages = query.order_by(desc(Message.created_at)).all()
        else:
            # When no conversation specified (legacy), limit to last 20
            messages = query.order_by(desc(Message.created_at)).limit(20).all()
        
        # Reverse to chronological order
        messages = list(reversed(messages))
        
        # Format messages for frontend
        formatted_messages = []
        for msg in messages:
            if msg.user_message:
                formatted_messages.append({
                    "role": "user",
                    "content": msg.user_message
                })
            if msg.bot_response:
                formatted_messages.append({
                    "role": "assistant",
                    "content": msg.bot_response
                })
        
        logger.info(f"Loaded {len(formatted_messages)} web messages for user {user.telegram_id}")
        return jsonify({"messages": formatted_messages})
    except Exception as e:
        logger.error(f"Error fetching messages: {str(e)}")
        return jsonify({
            "error": "Internal server error"
        }), 500

@app.route('/api/conversations', methods=['GET'])
@require_api_key
def get_conversations(**kwargs):
    user = kwargs['user']
    """Get all conversations for authenticated user"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Service temporarily unavailable"
        }), 503
    
    try:
        # Get all conversations with message counts in a single query (eliminates N+1 problem)
        from sqlalchemy import desc, func
        conversations_with_counts = db.session.query(
            Conversation,
            func.count(Message.id).label('message_count')
        ).outerjoin(Message).filter(
            Conversation.user_id == user.id
        ).group_by(Conversation.id).order_by(
            desc(Conversation.updated_at)
        ).all()
        
        # Format conversations for frontend
        conversation_list = []
        for conv, msg_count in conversations_with_counts:
            conversation_list.append({
                "id": conv.id,
                "title": conv.title,
                "created_at": conv.created_at.isoformat(),
                "updated_at": conv.updated_at.isoformat(),
                "message_count": msg_count
            })
        
        logger.info(f"Loaded {len(conversation_list)} conversations for user {user.telegram_id}")
        return jsonify({"conversations": conversation_list})
    except Exception as e:
        logger.error(f"Error fetching conversations: {str(e)}")
        return jsonify({
            "error": "Internal server error"
        }), 500

@app.route('/api/conversations', methods=['POST'])
@require_api_key
def create_conversation(**kwargs):
    user = kwargs['user']
    """Create a new conversation for authenticated user"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Service temporarily unavailable"
        }), 503
    
    try:
        # Get and validate title
        data = request.get_json() or {}
        title = data.get('title', 'New Chat').strip()[:255]  # Max length from schema
        if not title:
            title = 'New Chat'
        
        # Create new conversation
        new_conversation = Conversation(
            user_id=user.id,
            title=title
        )
        
        db.session.add(new_conversation)
        db.session.commit()
        
        logger.info(f"Created new conversation {new_conversation.id} for user {user.telegram_id}")
        
        return jsonify({
            "id": new_conversation.id,
            "title": new_conversation.title,
            "created_at": new_conversation.created_at.isoformat(),
            "updated_at": new_conversation.updated_at.isoformat(),
            "message_count": 0
        }), 201
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating conversation: {str(e)}")
        return jsonify({
            "error": "Internal server error"
        }), 500

@app.route('/api/conversations/<int:conversation_id>', methods=['DELETE'])
@require_api_key
def delete_conversation(conversation_id, **kwargs):
    user = kwargs['user']
    """Delete a conversation and all its messages"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Service temporarily unavailable"
        }), 503
    
    try:
        # Find the conversation
        conversation = Conversation.query.filter_by(
            id=conversation_id,
            user_id=user.id
        ).first()
        
        if not conversation:
            return jsonify({
                "error": "Conversation not found"
            }), 404
        
        # Delete the conversation (cascade will delete all messages)
        db.session.delete(conversation)
        db.session.commit()
        
        logger.info(f"Deleted conversation {conversation_id} for user {user.telegram_id}")
        
        return jsonify({
            "success": True,
            "message": "Conversation deleted"
        })
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting conversation: {str(e)}")
        return jsonify({
            "error": "Internal server error"
        }), 500

@app.route('/stats')
def stats():
    """Endpoint to view basic statistics about the bot usage with enhanced lock monitoring"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Database not available",
            "message": "Statistics require database connection"
        }), 503
    
    try:
        user_count = User.query.count()
        message_count = Message.query.count()
        recent_messages = Message.query.order_by(Message.created_at.desc()).limit(5).all()
        
        # Enhanced lock monitoring
        from datetime import timedelta
        current_time = datetime.utcnow()
        
        # Get ALL users with locks (active + stuck)
        all_locked_users = User.query.filter(User.processing_since.isnot(None)).all()
        locked_users_count = len(all_locked_users)
        
        # Define thresholds
        normal_threshold = timedelta(seconds=60)  # Lock should complete within 60s
        stuck_threshold = timedelta(minutes=5)  # Stuck after 5 minutes
        
        # Categorize locks
        active_locks = []  # Within normal range
        warning_locks = []  # Between 60s and 5min
        stuck_locks = []  # Over 5 minutes
        
        lock_ages = []
        for user in all_locked_users:
            lock_age = current_time - user.processing_since
            lock_ages.append(lock_age.total_seconds())
            
            lock_info = {
                "telegram_id": user.telegram_id,
                "username": user.username,
                "lock_age_seconds": int(lock_age.total_seconds()),
                "lock_age_minutes": round(lock_age.total_seconds() / 60, 1),
                "locked_since": user.processing_since.strftime("%Y-%m-%d %H:%M:%S")
            }
            
            if lock_age > stuck_threshold:
                lock_info["status"] = "🔴 STUCK"
                stuck_locks.append(lock_info)
            elif lock_age > normal_threshold:
                lock_info["status"] = "🟡 WARNING"
                warning_locks.append(lock_info)
            else:
                lock_info["status"] = "🟢 ACTIVE"
                active_locks.append(lock_info)
        
        # Calculate lock statistics
        lock_stats = {
            "total_locks": locked_users_count,
            "active_locks": len(active_locks),
            "warning_locks": len(warning_locks),
            "stuck_locks": len(stuck_locks),
            "oldest_lock_seconds": int(max(lock_ages)) if lock_ages else 0,
            "oldest_lock_minutes": round(max(lock_ages) / 60, 1) if lock_ages else 0,
            "average_lock_seconds": int(sum(lock_ages) / len(lock_ages)) if lock_ages else 0,
            "health_status": "🟢 HEALTHY" if len(stuck_locks) == 0 else "🔴 NEEDS ATTENTION"
        }
        
        # Format recent messages for display
        recent_msgs_formatted = []
        for msg in recent_messages:
            user = User.query.get(msg.user_id)
            if user:
                username = user.username or user.first_name or f"User {user.telegram_id}"
                recent_msgs_formatted.append({
                    "id": msg.id,
                    "user": username,
                    "message": msg.user_message[:30] + "..." if len(msg.user_message) > 30 else msg.user_message,
                    "time": msg.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                    "model": msg.model_used
                })
        
        return jsonify({
            "stats": {
                "total_users": user_count,
                "total_messages": message_count,
                "lock_health": lock_stats
            },
            "locks": {
                "active": active_locks[:5],  # Show first 5
                "warnings": warning_locks,
                "stuck": stuck_locks
            },
            "recent_messages": recent_msgs_formatted,
            "monitoring": {
                "thresholds": {
                    "normal": "< 60 seconds",
                    "warning": "60s - 5min",
                    "stuck": "> 5 minutes (auto-cleanup)"
                },
                "cleanup_info": "Stuck locks are automatically cleared every 5 minutes by background thread"
            }
        })
    except Exception as e:
        logger.error(f"Error generating stats: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/admin/clear_locks', methods=['POST'])
def admin_clear_locks():
    """Manually clear all stuck processing locks (ADMIN ONLY)"""
    admin_token = request.args.get('admin_token')
    if request.is_json and request.json:
        admin_token = admin_token or request.json.get('admin_token')
    expected_token = os.environ.get('ADMIN_EXPORT_TOKEN')
    
    if not expected_token:
        return jsonify({
            "error": "Admin functionality not configured",
            "message": "Admin token not set in environment"
        }), 503
    
    if not admin_token or admin_token != expected_token:
        logger.warning(f"Unauthorized lock cleanup attempt from {request.remote_addr}")
        return jsonify({
            "error": "Unauthorized",
            "message": "Valid admin token required"
        }), 401
    
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Database not available",
            "message": "Lock cleanup requires database connection"
        }), 503
    
    try:
        cleared_count = cleanup_stuck_processing_locks()
        return jsonify({
            "success": True,
            "message": f"Cleared {cleared_count} stuck processing locks",
            "cleared_count": cleared_count
        })
    except Exception as e:
        logger.error(f"Error in admin lock cleanup: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/export/conversations', methods=['GET'])
def export_conversations():
    """Export all conversations to a DOCX file (ADMIN ONLY - requires authentication)"""
    # CRITICAL SECURITY: Require admin authentication
    admin_token = request.args.get('admin_token')
    expected_token = os.environ.get('ADMIN_EXPORT_TOKEN')
    
    if not expected_token:
        return jsonify({
            "error": "Export feature not configured",
            "message": "Admin token not set in environment"
        }), 503
    
    if not admin_token or admin_token != expected_token:
        logger.warning(f"Unauthorized export attempt from {request.remote_addr}")
        return jsonify({
            "error": "Unauthorized",
            "message": "Valid admin token required"
        }), 401
    
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Database not available",
            "message": "Export requires database connection"
        }), 503
    
    try:
        # Query all messages with eager loading to prevent N+1 queries
        from sqlalchemy.orm import joinedload
        messages = Message.query.options(joinedload(Message.user)).order_by(Message.user_id, Message.created_at).all()
        
        if not messages:
            return jsonify({"error": "No conversations found"}), 404
        
        # Create DOCX document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Telegram Bot Conversations Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add export date
        export_info = doc.add_paragraph()
        export_info.add_run(f'Exported on: {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")}')
        export_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Blank line
        
        # Group messages by user
        current_user_id = None
        
        for msg in messages:
            # Use eager-loaded user from relationship (no additional query)
            user = msg.user
            
            # Add user header when switching to a new user
            if msg.user_id != current_user_id:
                current_user_id = msg.user_id
                
                # Add separator
                if doc.paragraphs[-1].text != '':
                    doc.add_page_break()
                
                # User info header with defensive NULL handling
                username = user.username if user and user.username else f"User {user.telegram_id if user else 'Unknown'}"
                user_heading = doc.add_heading(f'User: {username}', level=1)
                
                if user:
                    user_info = doc.add_paragraph()
                    user_info.add_run(f'Telegram ID: {user.telegram_id}').bold = True
                    
                    # Defensive handling for optional fields
                    name_parts = []
                    if user.first_name:
                        name_parts.append(user.first_name)
                    if user.last_name:
                        name_parts.append(user.last_name)
                    if name_parts:
                        user_info.add_run(f'\nName: {" ".join(name_parts)}')
                    
                    if user.registered_at:
                        user_info.add_run(f'\nRegistered: {user.registered_at.strftime("%Y-%m-%d %H:%M:%S")}')
                    
                    if user.credits is not None:
                        user_info.add_run(f'\nCredits: {user.credits}')
                
                doc.add_paragraph()  # Blank line
            
            # Add message timestamp
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f'[{msg.created_at.strftime("%Y-%m-%d %H:%M:%S")}]')
            timestamp_run.font.size = Pt(9)
            timestamp_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add user message
            user_msg_para = doc.add_paragraph()
            user_label = user_msg_para.add_run('User: ')
            user_label.bold = True
            user_label.font.color.rgb = RGBColor(0, 102, 204)
            user_msg_para.add_run(msg.user_message)
            
            # Add bot response
            if msg.bot_response:
                bot_msg_para = doc.add_paragraph()
                bot_label = bot_msg_para.add_run('Bot: ')
                bot_label.bold = True
                bot_label.font.color.rgb = RGBColor(34, 139, 34)
                bot_msg_para.add_run(msg.bot_response)
            
            # Add model info
            if msg.model_used:
                model_para = doc.add_paragraph()
                model_run = model_para.add_run(f'Model: {msg.model_used} | Credits: {msg.credits_charged}')
                model_run.font.size = Pt(8)
                model_run.font.color.rgb = RGBColor(128, 128, 128)
                model_run.italic = True
            
            doc.add_paragraph()  # Blank line between messages
        
        # Save document to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename with timestamp
        filename = f'telegram_conversations_{datetime.utcnow().strftime("%Y%m%d_%H%M%S")}.docx'
        
        logger.info(f"Exported {len(messages)} messages to DOCX")
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error exporting conversations: {str(e)}")
        return jsonify({"error": str(e)}), 500

def cleanup_stuck_processing_locks():
    """Clean up processing locks that are older than 5 minutes (stuck locks)"""
    if not DB_AVAILABLE:
        logger.warning("Cannot cleanup stuck locks - database not available")
        return 0
    
    try:
        from datetime import timedelta
        threshold = datetime.utcnow() - timedelta(minutes=5)
        
        stuck_users = User.query.filter(
            User.processing_since.isnot(None),
            User.processing_since < threshold
        ).all()
        
        count = len(stuck_users)
        
        if count > 0:
            logger.warning(f"Found {count} users with stuck processing locks (older than 5 minutes)")
            
            for user in stuck_users:
                lock_age = datetime.utcnow() - user.processing_since
                logger.info(f"Clearing stuck lock for user {user.telegram_id} ({user.username}) - lock age: {lock_age}")
                user.processing_since = None
            
            db.session.commit()
            logger.info(f"✓ Cleared {count} stuck processing locks")
        else:
            logger.info("No stuck processing locks found")
        
        return count
    except Exception as e:
        logger.error(f"Error cleaning up stuck locks: {str(e)}")
        db.session.rollback()
        return 0

def register_telegram_commands():
    """Register bot commands with Telegram so they appear in the command menu"""
    if not BOT_TOKEN:
        logger.warning("Cannot register commands - BOT_TOKEN not configured")
        return
    
    try:
        commands = [
            {"command": "start", "description": "Start the bot"},
            {"command": "help", "description": "Show help message"},
            {"command": "model", "description": "Switch AI model (DeepSeek/GPT-4o)"},
            {"command": "balance", "description": "Check credit balance"},
            {"command": "daily", "description": "Claim 25 free daily credits"},
            {"command": "buy", "description": "Purchase more credits"},
            {"command": "getapikey", "description": "Get API key for web chat access"},
            {"command": "clear", "description": "Clear conversation history"},
            {"command": "imagine", "description": "Generate photorealistic image (10 credits)"},
            {"command": "uncensored", "description": "Fully uncensored image generation (10 credits)"},
            {"command": "grok", "description": "Generate stylized image (8 credits)"},
            {"command": "edit", "description": "Image generation with text support (8 credits)"},
            {"command": "vid", "description": "Send photo + /vid caption for WAN 2.2 video"},
            {"command": "write", "description": "Professional writing mode (2 credits)"},
            {"command": "stats", "description": "System statistics (admin only)"}
        ]
        
        telegram_url = f"https://api.telegram.org/bot{BOT_TOKEN}/setMyCommands"
        response = requests.post(telegram_url, json={"commands": commands}, timeout=10)
        response_data = response.json()
        
        if response_data.get('ok'):
            logger.info(f"✓ Bot commands registered successfully ({len(commands)} commands)")
        else:
            logger.error(f"✗ Failed to register bot commands: {response_data}")
    except Exception as e:
        logger.error(f"Error registering commands: {str(e)}")

def register_telegram_webhook():
    """Automatically register Telegram webhook on app startup"""
    if not BOT_TOKEN:
        logger.warning("Cannot register webhook - BOT_TOKEN not configured")
        return
    
    try:
        # Use production domain (ko2bot.com) - locked in for production deployment
        domain = "ko2bot.com"
        
        # Build webhook URL
        webhook_url = f"https://{domain}/{BOT_TOKEN}"
        
        # Redact BOT_TOKEN from webhook_url for logging (do this early)
        safe_webhook_url = webhook_url.replace(BOT_TOKEN, "[REDACTED]")
        
        # Call Telegram API to set webhook
        telegram_url = f"https://api.telegram.org/bot{BOT_TOKEN}/setWebhook?url={webhook_url}"
        response = requests.get(telegram_url, timeout=10)
        response_data = response.json()
        
        if response_data.get('ok'):
            logger.info(f"✓ Webhook registered successfully: {safe_webhook_url}")
        else:
            logger.error(f"✗ Failed to register webhook to {safe_webhook_url}: {response_data}")
    except Exception as e:
        # Redact BOT_TOKEN from exception message before logging
        error_msg = str(e).replace(BOT_TOKEN, "[REDACTED]") if BOT_TOKEN else str(e)
        logger.error(f"Error registering webhook: {error_msg}")

# Clean up stuck processing locks on startup (within app context)
if DATABASE_URL and DB_AVAILABLE:
    with app.app_context():
        cleanup_stuck_processing_locks()

# Note: Webhook and commands registration moved to gunicorn.conf.py
# to avoid multiple workers registering simultaneously (rate limiting)

# Periodic lock cleanup function
def periodic_lock_cleanup():
    """Background thread to clean up stuck processing locks every 5 minutes"""
    while True:
        try:
            # Sleep first (5 minutes = 300 seconds) so startup cleanup runs first
            time.sleep(300)
            
            if DATABASE_URL and DB_AVAILABLE:
                with app.app_context():
                    cleared = cleanup_stuck_processing_locks()
                    if cleared > 0:
                        logger.warning(f"Periodic cleanup: Cleared {cleared} stuck locks")
                    else:
                        logger.debug("Periodic cleanup: No stuck locks found")
            else:
                logger.debug("Periodic cleanup: Skipping - database not available")
        except Exception as e:
            logger.error(f"Error in periodic lock cleanup: {str(e)}")

# Keepalive function
def keep_alive():
    """Function to ping the app every 4 minutes to prevent Replit from sleeping"""
    while True:
        try:
            logger.info("Pinging server to keep it alive...")
            requests.get(KEEPALIVE_URL, timeout=10)
            logger.info("Ping successful")
        except Exception as e:
            logger.error(f"Error pinging server: {str(e)}")
        # Sleep for 4 minutes (240 seconds)
        time.sleep(240)

# Start background threads
keepalive_thread = threading.Thread(target=keep_alive, daemon=True)
keepalive_thread.start()
logger.info("Started keepalive thread")

if DATABASE_URL and DB_AVAILABLE:
    lock_cleanup_thread = threading.Thread(target=periodic_lock_cleanup, daemon=True)
    lock_cleanup_thread.start()
    logger.info("Started periodic lock cleanup thread (runs every 5 minutes)")

def verify_nowpayments_ipn(ipn_secret, raw_body_bytes, received_signature):
    """
    Verify NOWPayments IPN callback signature using HMAC-SHA512
    
    Args:
        ipn_secret (str): IPN secret key from NOWPayments
        raw_body_bytes (bytes): Raw request body as bytes (not re-serialized JSON)
        received_signature (str): Signature from x-nowpayments-sig header
    
    Returns:
        bool: True if signature is valid, False otherwise
    """
    try:
        # Calculate HMAC-SHA512 signature on the raw request body bytes
        # CRITICAL: Must use raw body exactly as received, not re-serialized JSON
        signature = hmac.new(
            ipn_secret.encode('utf-8'),
            raw_body_bytes,
            hashlib.sha512
        ).hexdigest()
        
        # Normalize both signatures (strip whitespace, lowercase) before comparison
        computed_sig = signature.strip().lower()
        received_sig = received_signature.strip().lower()
        
        # Use constant-time comparison to prevent timing attacks
        return hmac.compare_digest(computed_sig, received_sig)
    except Exception as e:
        logger.error(f"IPN signature verification error: {str(e)}")
        return False

@app.route('/v1/chat/completions', methods=['POST'])
def chat_completions_proxy():
    """OpenAI-compatible proxy endpoint with FULL jailbreak system (same as Telegram)
    
    This endpoint:
    - Authenticates users via API key
    - Fetches last 10 messages for conversation context
    - Injects jailbreak system prompt
    - Deducts credits (1-3 per request based on user's model preference)
    - Uses generate_response() with refusal detection
    - Returns OpenAI-compatible streaming responses
    - Stores messages with platform='web'
    """
    # Start timing for performance monitoring
    request_start_time = time.time()
    
    logger.info("=" * 80)
    logger.info("WEB CHAT REQUEST RECEIVED")
    logger.info(f"Request method: {request.method}")
    logger.info(f"Request path: {request.path}")
    logger.info("=" * 80)
    
    if not DB_AVAILABLE:
        return jsonify({
            "error": {
                "message": "Service temporarily unavailable - database not connected",
                "type": "service_unavailable",
                "code": "db_unavailable"
            }
        }), 503
    
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '):
        logger.warning(f"Missing or invalid Authorization header from {request.remote_addr}")
        return jsonify({
            "error": {
                "message": "Invalid API key. Get yours from Telegram: /getapikey",
                "type": "invalid_request_error",
                "code": "invalid_api_key"
            }
        }), 401
    
    api_key = auth_header[7:]
    
    try:
        user = User.query.filter_by(api_key=api_key).first()
        
        if not user:
            logger.warning(f"Invalid API key attempted from {request.remote_addr}")
            return jsonify({
                "error": {
                    "message": "Invalid API key. Get yours from Telegram: /getapikey",
                    "type": "invalid_request_error",
                    "code": "invalid_api_key"
                }
            }), 401
        
        logger.info(f"Authenticated web user: {user.telegram_id} (username: {user.username})")
        
        # Calculate credit cost based on user's preferred model (DeepSeek=1, GPT-4o=3)
        selected_model = user.preferred_model or 'deepseek/deepseek-chat-v3-0324'
        credits_to_deduct = 3 if 'gpt-4o' in selected_model.lower() or 'chatgpt' in selected_model.lower() else 1
        
        from telegram_handler import deduct_credits
        success, daily_used, purchased_used, credit_warning = deduct_credits(user, credits_to_deduct)
        
        if not success:
            total_credits = user.daily_credits + user.credits
            logger.warning(f"Insufficient credits for user {user.telegram_id}: {total_credits} credits")
            
            domain = os.environ.get("REPLIT_DOMAINS", "").split(',')[0] if os.environ.get("REPLIT_DOMAINS") else os.environ.get("REPLIT_DEV_DOMAIN")
            buy_url = f"https://{domain}/buy?telegram_id={user.telegram_id}" if domain else f"/buy?telegram_id={user.telegram_id}"
            
            return jsonify({
                "error": {
                    "message": f"Insufficient credits. Balance: {total_credits}. Purchase: {buy_url}",
                    "type": "insufficient_quota",
                    "code": "insufficient_credits",
                    "param": {
                        "balance": total_credits,
                        "buy_url": buy_url
                    }
                }
            }), 402
        
        db.session.commit()
        logger.info(f"Deducted {credits_to_deduct} credit(s) (daily: {daily_used}, purchased: {purchased_used}). New balance: daily={user.daily_credits}, purchased={user.credits}")
        
        payload = request.get_json()
        if not payload:
            return jsonify({
                "error": {
                    "message": "Invalid JSON payload",
                    "type": "invalid_request_error",
                    "code": "invalid_json"
                }
            }), 400
        
        messages = payload.get('messages', [])
        if not messages:
            return jsonify({
                "error": {
                    "message": "No messages provided",
                    "type": "invalid_request_error",
                    "code": "no_messages"
                }
            }), 400
        
        # Get user's current message (last message in the array)
        user_message = messages[-1].get('content', '') if messages else ''
        if not user_message:
            return jsonify({
                "error": {
                    "message": "Empty user message",
                    "type": "invalid_request_error",
                    "code": "empty_message"
                }
            }), 400
        
        logger.info(f"Web user message: {user_message[:100]}...")
        
        # Check for memory commands (! memorize, ! memories, ! forget)
        from memory_utils import parse_memory_command, store_memory, get_user_memories, delete_memory, format_memories_for_display
        
        command_type, command_data = parse_memory_command(user_message)
        
        if command_type == 'store':
            try:
                memory = store_memory(user.id, command_data, platform='web')
                response_content = f"✅ Memory saved! (ID: {memory.id})\n\n📝 {command_data}\n\n💡 Use `! memories` to view all saved memories."
                
                # Refund credit since this is a command, not an LLM request
                user.credits += purchased_used
                user.daily_credits += daily_used
                db.session.commit()
                
                return jsonify({
                    "id": f"chatcmpl-{int(time.time())}",
                    "object": "chat.completion",
                    "created": int(time.time()),
                    "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                    "choices": [{
                        "index": 0,
                        "message": {
                            "role": "assistant",
                            "content": response_content
                        },
                        "finish_reason": "stop"
                    }],
                    "usage": {
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0
                    }
                })
            except Exception as e:
                logger.error(f"Failed to store memory: {e}")
                return jsonify({
                    "error": {
                        "message": "Failed to save memory. Please try again.",
                        "type": "server_error",
                        "code": "memory_storage_failed"
                    }
                }), 500
        
        elif command_type == 'list':
            try:
                memories = get_user_memories(user.id)
                response_content = format_memories_for_display(memories)
                
                # Refund credit since this is a command, not an LLM request
                user.credits += purchased_used
                user.daily_credits += daily_used
                db.session.commit()
                
                return jsonify({
                    "id": f"chatcmpl-{int(time.time())}",
                    "object": "chat.completion",
                    "created": int(time.time()),
                    "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                    "choices": [{
                        "index": 0,
                        "message": {
                            "role": "assistant",
                            "content": response_content
                        },
                        "finish_reason": "stop"
                    }],
                    "usage": {
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0
                    }
                })
            except Exception as e:
                logger.error(f"Failed to list memories: {e}")
                return jsonify({
                    "error": {
                        "message": "Failed to retrieve memories. Please try again.",
                        "type": "server_error",
                        "code": "memory_retrieval_failed"
                    }
                }), 500
        
        elif command_type == 'forget':
            try:
                memory_id = command_data
                success = delete_memory(user.id, memory_id)
                if success:
                    response_content = f"🗑️ Memory [{memory_id}] deleted successfully."
                else:
                    response_content = f"❌ Memory [{memory_id}] not found or doesn't belong to you."
                
                # Refund credit since this is a command, not an LLM request
                user.credits += purchased_used
                user.daily_credits += daily_used
                db.session.commit()
                
                return jsonify({
                    "id": f"chatcmpl-{int(time.time())}",
                    "object": "chat.completion",
                    "created": int(time.time()),
                    "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                    "choices": [{
                        "index": 0,
                        "message": {
                            "role": "assistant",
                            "content": response_content
                        },
                        "finish_reason": "stop"
                    }],
                    "usage": {
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0
                    }
                })
            except Exception as e:
                logger.error(f"Failed to delete memory: {e}")
                return jsonify({
                    "error": {
                        "message": "Failed to delete memory. Please try again.",
                        "type": "server_error",
                        "code": "memory_deletion_failed"
                    }
                }), 500
        
        # Get conversation_id from payload, or create/get a default conversation
        conversation_id = payload.get('conversation_id')
        
        if conversation_id:
            # Verify conversation belongs to user
            conversation = Conversation.query.filter_by(
                id=conversation_id,
                user_id=user.id
            ).first()
            
            if not conversation:
                logger.warning(f"Invalid conversation_id {conversation_id} for user {user.telegram_id}")
                return jsonify({
                    "error": {
                        "message": "Conversation not found",
                        "type": "invalid_request_error",
                        "code": "invalid_conversation"
                    }
                }), 404
        else:
            # Auto-create or get default conversation for this user
            conversation = Conversation.query.filter_by(
                user_id=user.id
            ).order_by(desc(Conversation.updated_at)).first()
            
            if not conversation:
                # Create first conversation for new user
                conversation = Conversation(
                    user_id=user.id,
                    title='New Chat'
                )
                db.session.add(conversation)
                db.session.commit()
                logger.info(f"Created first conversation {conversation.id} for user {user.telegram_id}")
            
            conversation_id = conversation.id
        
        logger.info(f"Using conversation {conversation_id}")
        
        # Load last 10 messages efficiently (optimized single query)
        recent_messages = Message.query.filter_by(
            conversation_id=conversation_id,
            platform='web'
        ).order_by(Message.created_at.desc()).limit(10).all()
        
        # Reverse to get chronological order (oldest first)
        recent_messages.reverse()
        
        logger.info(f"Loaded {len(recent_messages)} previous web messages for context (conversation {conversation_id})")
        
        # Build conversation history (same format as Telegram)
        conversation_history = []
        for msg in recent_messages:
            if msg.user_message:
                conversation_history.append({
                    "role": "user",
                    "content": msg.user_message
                })
            if msg.bot_response:
                conversation_history.append({
                    "role": "assistant",
                    "content": msg.bot_response
                })
        
        logger.info(f"Built conversation history with {len(conversation_history)} messages")
        
        # Use the SAME generate_response function as Telegram (includes jailbreak, refusal detection, etc.)
        is_streaming = payload.get('stream', True)
        
        if is_streaming:
            # For streaming, we need to capture the response and convert to OpenAI format
            from flask import Response
            
            accumulated_response = ""
            
            def update_callback(chunk):
                nonlocal accumulated_response
                accumulated_response += chunk
            
            # Call generate_response with streaming callback (include user_id for memory injection and user's preferred model)
            selected_model = user.preferred_model or 'deepseek/deepseek-chat-v3-0324'
            bot_response = generate_response(
                user_message=user_message,
                conversation_history=conversation_history,
                use_streaming=True,
                update_callback=update_callback,
                writing_mode=False,
                user_id=user.id,
                model=selected_model
            )
            
            # Auto-title if first message
            auto_title_conversation_if_first_message(conversation, user_message, conversation_id)
            
            # Store message in database with conversation association
            message_record = Message(
                user_id=user.id,
                conversation_id=conversation_id,
                user_message=user_message[:1000],
                bot_response=bot_response[:10000] if bot_response else "",
                model_used=payload.get('model', 'openai/chatgpt-4o-latest'),
                credits_charged=1,
                platform='web'
            )
            db.session.add(message_record)
            
            # Update conversation's updated_at timestamp
            conversation.updated_at = datetime.utcnow()
            
            db.session.commit()
            message_id = message_record.id
            
            # Create transaction record
            transaction = Transaction(
                user_id=user.id,
                credits_used=1,
                message_id=message_id,
                transaction_type='web_message',
                description=f"Web chat message"
            )
            db.session.add(transaction)
            db.session.commit()
            
            logger.info(f"Stored web message (id={message_id}) and transaction")
            
            # Convert to OpenAI streaming format
            def generate_openai_stream():
                # Send chunks
                words = bot_response.split(' ') if bot_response else []
                for i, word in enumerate(words):
                    chunk_data = {
                        "id": f"chatcmpl-{uuid.uuid4()}",
                        "object": "chat.completion.chunk",
                        "created": int(time.time()),
                        "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                        "choices": [{
                            "index": 0,
                            "delta": {"content": word + (' ' if i < len(words) - 1 else '')},
                            "finish_reason": None
                        }]
                    }
                    yield f"data: {json.dumps(chunk_data)}\n\n"
                
                # Send final chunk
                final_chunk = {
                    "id": f"chatcmpl-{uuid.uuid4()}",
                    "object": "chat.completion.chunk",
                    "created": int(time.time()),
                    "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                    "choices": [{
                        "index": 0,
                        "delta": {},
                        "finish_reason": "stop"
                    }]
                }
                yield f"data: {json.dumps(final_chunk)}\n\n"
                yield "data: [DONE]\n\n"
            
            return Response(generate_openai_stream(), content_type='text/event-stream')
        else:
            # Non-streaming response (include user_id for memory injection)
            bot_response = generate_response(
                user_message=user_message,
                conversation_history=conversation_history,
                use_streaming=False,
                writing_mode=False,
                user_id=user.id
            )
            
            # Auto-title if first message
            auto_title_conversation_if_first_message(conversation, user_message, conversation_id)
            
            # Store message with conversation association
            message_record = Message(
                user_id=user.id,
                conversation_id=conversation_id,
                user_message=user_message[:1000],
                bot_response=bot_response[:10000] if bot_response else "",
                model_used=payload.get('model', 'openai/chatgpt-4o-latest'),
                credits_charged=1,
                platform='web'
            )
            db.session.add(message_record)
            
            # Update conversation's updated_at timestamp
            conversation.updated_at = datetime.utcnow()
            
            db.session.commit()
            
            # Create transaction
            transaction = Transaction(
                user_id=user.id,
                credits_used=1,
                message_id=message_record.id,
                transaction_type='web_message',
                description=f"Web chat message"
            )
            db.session.add(transaction)
            db.session.commit()
            
            # Log request timing
            request_time_ms = (time.time() - request_start_time) * 1000
            logger.info(f"Web chat request completed in {request_time_ms:.2f}ms (non-streaming)")
            
            # Return OpenAI-compatible format
            return jsonify({
                "id": f"chatcmpl-{uuid.uuid4()}",
                "object": "chat.completion",
                "created": int(time.time()),
                "model": payload.get('model', 'openai/chatgpt-4o-latest'),
                "choices": [{
                    "index": 0,
                    "message": {
                        "role": "assistant",
                        "content": bot_response
                    },
                    "finish_reason": "stop"
                }],
                "usage": {
                    "prompt_tokens": 0,
                    "completion_tokens": 0,
                    "total_tokens": 0
                }
            })
            
    except Exception as e:
        # Log request timing even on error
        request_time_ms = (time.time() - request_start_time) * 1000
        logger.error(f"Web chat request failed after {request_time_ms:.2f}ms: {str(e)}", exc_info=True)
        
        # Refund credits on error (only if they were deducted)
        try:
            if 'user' in locals() and 'purchased_used' in locals() and 'daily_used' in locals():
                user.credits += purchased_used
                user.daily_credits += daily_used
                db.session.commit()
                logger.info(f"Refunded {purchased_used + daily_used} credits due to error")
        except:
            pass
        
        return jsonify({
            "error": {
                "message": f"Internal server error: {str(e)}",
                "type": "server_error",
                "code": "internal_error"
            }
        }), 500

@app.route(f'/{BOT_TOKEN}', methods=['POST'])
def webhook():
    """Webhook endpoint to receive updates from Telegram
    
    CRITICAL: Always returns 200 OK to Telegram to prevent webhook delivery issues.
    Errors are handled internally and logged, but never cause 5xx responses.
    """
    logger.info("=" * 80)
    logger.info("WEBHOOK REQUEST RECEIVED")
    logger.info(f"Request method: {request.method}")
    logger.info(f"Request path: {request.path}")
    logger.info(f"Request headers: {dict(request.headers)}")
    logger.info("=" * 80)
    
    if not BOT_TOKEN:
        logger.error("Bot token not configured!")
        # Still return 200 to Telegram to prevent delivery issues
        return jsonify({"status": "error", "message": "Bot not configured"}), 200
    
    # CRITICAL: Capture request data BEFORE starting background thread
    # (request context is not available inside the thread)
    try:
        update = request.get_json()
        logger.info(f"Received Telegram update: {json.dumps(update, indent=2)}")
    except Exception as e:
        logger.error(f"Failed to parse webhook JSON: {str(e)}")
        # Still return 200 to prevent Telegram from retrying bad data
        return jsonify({"status": "error", "message": "Invalid JSON"}), 200
    
    # CRITICAL: Process update in background thread to prevent timeouts
    # and always return 200 OK to Telegram immediately
    def process_in_background():
        try:
            # CRITICAL: Wrap with app context for database operations
            with app.app_context():
                logger.info("Calling process_update()...")
                process_update(update)
                logger.info("process_update() completed successfully")
            
        except Exception as e:
            logger.error(f"Error processing webhook in background: {str(e)}", exc_info=True)
            # Try to notify user of the error if we can extract chat_id
            try:
                chat_id = None
                if update.get('message'):
                    chat_id = update['message'].get('chat', {}).get('id')
                elif update.get('callback_query'):
                    chat_id = update['callback_query'].get('message', {}).get('chat', {}).get('id')
                
                if chat_id:
                    send_message(chat_id, f"⚠️ An error occurred processing your request. Please try again or contact support.\n\nError: {str(e)[:100]}")
            except:
                logger.error("Could not send error notification to user")
    
    # Start background processing
    thread = threading.Thread(target=process_in_background)
    thread.daemon = True
    thread.start()
    
    # ALWAYS return 200 OK immediately to Telegram
    return jsonify({"status": "ok"}), 200

@app.route('/set_webhook', methods=['GET'])
def set_webhook():
    """Helper endpoint to set the webhook URL for Telegram"""
    if not BOT_TOKEN:
        return jsonify({"error": "Bot token not configured"}), 500
        
    try:
        # Get domain from environment or request parameter
        url = request.args.get('url')
        if not url:
            # Try to auto-detect domain from Replit environment
            domain = os.environ.get("REPLIT_DOMAINS", "").split(',')[0] if os.environ.get("REPLIT_DOMAINS") else os.environ.get("REPLIT_DEV_DOMAIN")
            if domain:
                if not domain.startswith('http'):
                    url = f"https://{domain}"
                else:
                    url = domain
            else:
                return jsonify({"error": "URL parameter is required or REPLIT_DOMAINS must be set"}), 400
            
        webhook_url = f"{url}/{BOT_TOKEN}"
        telegram_url = f"https://api.telegram.org/bot{BOT_TOKEN}/setWebhook?url={webhook_url}"
        
        # Make the actual request to Telegram
        response = requests.get(telegram_url)
        response_data = response.json()
        
        if response_data.get('ok'):
            logger.info(f"Webhook set successfully to: {webhook_url}")
            return jsonify({
                "success": True,
                "message": "Webhook set successfully",
                "webhook_url": webhook_url,
                "telegram_response": response_data
            })
        else:
            logger.error(f"Failed to set webhook: {response_data}")
            return jsonify({
                "success": False,
                "error": "Failed to set webhook",
                "telegram_response": response_data
            }), 400
    except Exception as e:
        logger.error(f"Error setting webhook: {str(e)}")
        return jsonify({"error": str(e)}), 500
        
@app.route('/test', methods=['POST'])
def test_bot():
    """Test endpoint to simulate a message to the bot without Telegram"""
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return jsonify({"error": "Request must include 'message' field"}), 400
            
        user_message = data['message']
        telegram_id = data.get('telegram_id', 1234567890)  # Allow custom telegram_id, default to test ID
        
        # Create a mock Telegram update
        mock_update = {
            "update_id": 123456789,
            "message": {
                "message_id": 100,
                "from": {
                    "id": telegram_id,
                    "is_bot": False,
                    "first_name": "Test",
                    "username": "testuser"
                },
                "chat": {
                    "id": telegram_id,  # Same as from.id
                    "first_name": "Test",
                    "username": "testuser",
                    "type": "private"
                },
                "date": int(time.time()),
                "text": user_message
            }
        }
        
        # Process the update just like a real Telegram update
        process_update(mock_update)
        
        # Get the response (process_update handles sending, but we also return it)
        response = generate_response(user_message)
        
        return jsonify({
            "status": "success",
            "user_message": user_message,
            "bot_response": response
        })
    except Exception as e:
        logger.error(f"Error in test endpoint: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/buy', methods=['GET'])
def buy_credits_page():
    """Buy credits page - displays credit packages and payment options"""
    telegram_id = request.args.get('telegram_id', '')
    
    if not telegram_id:
        return "Error: telegram_id parameter is required", 400
    
    html = f'''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Buy Credits</title>
        <style>
            * {{ box-sizing: border-box; margin: 0; padding: 0; }}
            body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
                   background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                   min-height: 100vh; padding: 20px; }}
            .container {{ max-width: 600px; margin: 0 auto; background: white; 
                         border-radius: 20px; padding: 30px; box-shadow: 0 20px 60px rgba(0,0,0,0.3); }}
            h1 {{ color: #333; margin-bottom: 10px; text-align: center; }}
            .subtitle {{ color: #666; text-align: center; margin-bottom: 30px; font-size: 14px; }}
            .packages {{ display: grid; gap: 15px; margin-bottom: 25px; }}
            .package {{ border: 2px solid #e0e0e0; border-radius: 12px; padding: 20px; 
                       cursor: pointer; transition: all 0.3s; position: relative; }}
            .package:hover {{ border-color: #667eea; transform: translateY(-2px); 
                            box-shadow: 0 5px 15px rgba(102,126,234,0.2); }}
            .package.selected {{ border-color: #667eea; background: #f8f9ff; }}
            .package-title {{ font-size: 20px; font-weight: bold; color: #333; margin-bottom: 5px; }}
            .package-price {{ font-size: 28px; font-weight: bold; color: #667eea; margin-bottom: 5px; }}
            .package-desc {{ color: #666; font-size: 14px; }}
            .form-group {{ margin-bottom: 20px; }}
            label {{ display: block; margin-bottom: 8px; font-weight: 600; color: #333; }}
            select {{ width: 100%; padding: 12px; border: 2px solid #e0e0e0; border-radius: 8px; 
                     font-size: 16px; transition: border-color 0.3s; }}
            select:focus {{ outline: none; border-color: #667eea; }}
            button {{ width: 100%; padding: 15px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                     color: white; border: none; border-radius: 8px; font-size: 18px; font-weight: bold; 
                     cursor: pointer; transition: transform 0.2s; }}
            button:hover {{ transform: translateY(-2px); }}
            button:disabled {{ opacity: 0.6; cursor: not-allowed; }}
            .result {{ margin-top: 20px; padding: 15px; border-radius: 8px; display: none; }}
            .result.success {{ background: #d4edda; border: 1px solid #c3e6cb; color: #155724; }}
            .result.error {{ background: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; }}
            .payment-info {{ margin-top: 15px; }}
            .payment-info div {{ margin-bottom: 10px; padding: 10px; background: #f8f9fa; 
                                border-radius: 5px; word-break: break-all; }}
            .payment-info strong {{ display: block; margin-bottom: 5px; color: #667eea; }}
            .loading {{ text-align: center; padding: 15px; }}
        </style>
    </head>
    <body>
        <div class="container">
            <h1>💎 Buy Credits</h1>
            <p class="subtitle">💰 Save 35% with crypto - no Telegram fees! Direct payment = better prices.</p>
            
            <div class="packages" id="packages">
                <div class="package" data-credits="200" data-price="5">
                    <div class="package-title">Starter Pack 🔥</div>
                    <div class="package-price">$5</div>
                    <div class="package-desc">200 credits • 35% cheaper than Telegram Stars</div>
                </div>
                <div class="package" data-credits="400" data-price="10">
                    <div class="package-title">Popular Pack 🔥</div>
                    <div class="package-price">$10</div>
                    <div class="package-desc">400 credits • 35% cheaper than Telegram Stars</div>
                </div>
                <div class="package" data-credits="800" data-price="20">
                    <div class="package-title">Value Pack 🔥</div>
                    <div class="package-price">$20</div>
                    <div class="package-desc">800 credits • 35% cheaper than Telegram Stars</div>
                </div>
                <div class="package" data-credits="2000" data-price="50">
                    <div class="package-title">Premium Pack 🔥</div>
                    <div class="package-price">$50</div>
                    <div class="package-desc">2,000 credits • 35% cheaper than Telegram Stars</div>
                </div>
            </div>
            
            <div class="form-group">
                <label for="currency">Payment Currency</label>
                <select id="currency">
                    <option value="">Loading currencies...</option>
                </select>
            </div>
            
            <button id="createPayment" onclick="createPayment()" disabled>Create Payment</button>
            
            <div id="result" class="result"></div>
        </div>
        
        <script>
            const telegramId = '{telegram_id}';
            let selectedPackage = null;
            
            // Package selection
            document.querySelectorAll('.package').forEach(pkg => {{
                pkg.addEventListener('click', () => {{
                    document.querySelectorAll('.package').forEach(p => p.classList.remove('selected'));
                    pkg.classList.add('selected');
                    selectedPackage = {{
                        credits: parseInt(pkg.dataset.credits),
                        price: parseFloat(pkg.dataset.price)
                    }};
                    updateButtonState();
                }});
            }});
            
            // Load currencies
            fetch('/api/crypto/currencies')
                .then(r => r.json())
                .then(data => {{
                    const select = document.getElementById('currency');
                    const currencies = data.currencies || [];
                    select.innerHTML = '<option value="">-- Select Currency --</option>' + 
                        currencies.map(c => `<option value="${{c}}">${{c.toUpperCase()}}</option>`).join('');
                    updateButtonState();
                }})
                .catch(err => {{
                    document.getElementById('currency').innerHTML = '<option value="">Error loading currencies</option>';
                }});
            
            function updateButtonState() {{
                const btn = document.getElementById('createPayment');
                const currency = document.getElementById('currency').value;
                btn.disabled = !selectedPackage || !currency;
            }}
            
            document.getElementById('currency').addEventListener('change', updateButtonState);
            
            async function createPayment() {{
                const resultDiv = document.getElementById('result');
                const button = document.getElementById('createPayment');
                const currency = document.getElementById('currency').value;
                
                if (!selectedPackage || !currency) {{
                    return;
                }}
                
                button.disabled = true;
                resultDiv.className = 'result';
                resultDiv.innerHTML = '<div class="loading">Creating payment...</div>';
                resultDiv.style.display = 'block';
                
                try {{
                    const response = await fetch('/api/crypto/create-payment', {{
                        method: 'POST',
                        headers: {{ 'Content-Type': 'application/json' }},
                        body: JSON.stringify({{
                            telegram_id: telegramId,
                            credits: selectedPackage.credits,
                            pay_currency: currency
                        }})
                    }});
                    
                    const data = await response.json();
                    
                    if (data.success) {{
                        resultDiv.className = 'result success';
                        resultDiv.innerHTML = `
                            <h3>✅ Payment Created Successfully!</h3>
                            <div class="payment-info">
                                <div><strong>Payment ID:</strong> ${{data.payment_id}}</div>
                                <div><strong>Send ${{data.pay_amount}} ${{data.pay_currency}}</strong> to this address:</div>
                                <div style="font-family: monospace; font-size: 14px;"><strong>Address:</strong> ${{data.pay_address}}</div>
                                <div><strong>Status:</strong> ${{data.payment_status}}</div>
                                <div style="margin-top: 15px; padding: 10px; background: #fff3cd; border: 1px solid #ffeaa7; border-radius: 5px;">
                                    ⏳ Credits will be added automatically when payment is confirmed
                                </div>
                            </div>
                        `;
                    }} else {{
                        resultDiv.className = 'result error';
                        resultDiv.innerHTML = `<strong>❌ Error:</strong> ${{data.error || 'Payment creation failed'}}`;
                        button.disabled = false;
                    }}
                }} catch (err) {{
                    resultDiv.className = 'result error';
                    resultDiv.innerHTML = `<strong>❌ Error:</strong> ${{err.message}}`;
                    button.disabled = false;
                }}
            }}
        </script>
    </body>
    </html>
    '''
    
    return render_template_string(html)

@app.route('/api/crypto/currencies', methods=['GET'])
def get_crypto_currencies():
    """Get list of available cryptocurrencies"""
    if not nowpayments:
        return jsonify({"error": "Crypto payments not configured"}), 503
    
    try:
        currencies_response = nowpayments.currencies()
        # Extract the currencies array from the response
        if isinstance(currencies_response, dict) and 'currencies' in currencies_response:
            currencies = currencies_response['currencies']
        else:
            currencies = currencies_response
        return jsonify({"currencies": currencies}), 200
    except Exception as e:
        logger.error(f"Error fetching crypto currencies: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/crypto/minimum-amount', methods=['GET'])
def get_minimum_payment_amount():
    """Get minimum payment amount for a cryptocurrency"""
    if not nowpayments:
        return jsonify({"error": "Crypto payments not configured"}), 503
    
    try:
        currency = request.args.get('currency', 'btc')
        min_amount_data = nowpayments.minimum_payment_amount(currency_from='usd', currency_to=currency.lower())
        return jsonify(min_amount_data), 200
    except Exception as e:
        logger.error(f"Error fetching minimum payment amount: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/crypto/create-payment', methods=['POST'])
def create_crypto_payment():
    """Create a crypto payment using NOWPayments"""
    if not nowpayments:
        return jsonify({"error": "Crypto payments not configured"}), 503
    
    if not DB_AVAILABLE:
        return jsonify({"error": "Database not available"}), 503
    
    try:
        data = request.get_json()
        if not data or 'credits' not in data or 'pay_currency' not in data:
            return jsonify({"error": "Request must include 'credits' and 'pay_currency' fields"}), 400
        
        credits = int(data['credits'])
        pay_currency = data['pay_currency']
        user_telegram_id = data.get('telegram_id')
        
        if credits <= 0:
            return jsonify({"error": "Credits must be a positive number"}), 400
        
        if not user_telegram_id:
            return jsonify({"error": "telegram_id is required"}), 400
        
        # Convert telegram_id to integer (Telegram IDs are always numeric)
        try:
            user_telegram_id = int(user_telegram_id)
        except (ValueError, TypeError):
            return jsonify({"error": "telegram_id must be a valid number"}), 400
        
        # Calculate amount in USD based on credit package (half price)
        CREDIT_PACKAGES = {
            200: 5.0,     # $5 → 200 credits (2.5¢/credit)
            400: 10.0,    # $10 → 400 credits (2.5¢/credit)
            800: 20.0,    # $20 → 800 credits (2.5¢/credit)
            2000: 50.0    # $50 → 2,000 credits (2.5¢/credit)
        }
        
        if credits not in CREDIT_PACKAGES:
            return jsonify({"error": f"Invalid credit amount. Valid options: {list(CREDIT_PACKAGES.keys())}"}), 400
        
        price_amount = CREDIT_PACKAGES[credits]
        
        # Check minimum payment amount
        try:
            min_amount_data = nowpayments.minimum_payment_amount(currency_from='usd', currency_to=pay_currency.lower())
            min_fiat_amount = float(min_amount_data.get('fiat_equivalent', 0))
            if price_amount < min_fiat_amount:
                # Find the smallest package that meets minimum
                suitable_packages = [c for c, p in CREDIT_PACKAGES.items() if p >= min_fiat_amount]
                min_package = min(suitable_packages) if suitable_packages else 200
                return jsonify({
                    "error": f"Amount too low. Minimum payment for {pay_currency.upper()} is ${min_fiat_amount:.2f}. Please select {min_package} credits package or higher."
                }), 400
        except Exception as e:
            logger.warning(f"Could not check minimum amount: {str(e)}")
        
        # Get or create user
        user = User.query.filter_by(telegram_id=user_telegram_id).first()
        if not user:
            logger.warning(f"User {user_telegram_id} not found, creating new user")
            user = User(telegram_id=user_telegram_id, credits=100)
            db.session.add(user)
            db.session.commit()
        
        # Generate unique order ID
        order_id = f"crypto_order_{user_telegram_id}_{int(datetime.utcnow().timestamp())}"
        
        # Get domain for IPN callback
        domain = os.environ.get("REPLIT_DOMAINS", "").split(',')[0] if os.environ.get("REPLIT_DOMAINS") else os.environ.get("REPLIT_DEV_DOMAIN")
        
        if not domain:
            logger.error("No domain configured for IPN callback")
            return jsonify({"error": "Domain not configured"}), 500
        
        if not domain.startswith('http'):
            domain = f"https://{domain}"
        
        logger.info(f"Creating crypto payment for {credits} credits (${price_amount:.2f}) in {pay_currency} for user {user_telegram_id}")
        
        # Create payment via NOWPayments
        payment_response = nowpayments.create_payment(
            price_amount=price_amount,
            price_currency='usd',
            pay_currency=pay_currency.lower(),
            ipn_callback_url=f"{domain}/api/crypto/ipn",
            order_id=order_id,
            order_description=f'Purchase {credits} credits for AI chat bot'
        )
        
        # Create CryptoPayment record in database
        crypto_payment = CryptoPayment(
            user_id=user.id,
            payment_id=payment_response['payment_id'],
            order_id=order_id,
            credits_purchased=credits,
            price_amount=price_amount,
            price_currency='USD',
            pay_amount=payment_response.get('pay_amount'),
            pay_currency=pay_currency.upper(),
            pay_address=payment_response.get('pay_address'),
            payment_status=payment_response.get('payment_status', 'waiting')
        )
        db.session.add(crypto_payment)
        db.session.commit()
        
        logger.info(f"Created crypto payment record {crypto_payment.id} with payment_id {payment_response['payment_id']}")
        
        return jsonify({
            "success": True,
            "payment_id": payment_response['payment_id'],
            "pay_address": payment_response['pay_address'],
            "pay_amount": payment_response['pay_amount'],
            "pay_currency": pay_currency.upper(),
            "payment_status": payment_response.get('payment_status'),
            "order_id": order_id
        }), 200
        
    except Exception as e:
        logger.error(f"Error creating crypto payment: {str(e)}")
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/crypto/ipn', methods=['POST'])
def crypto_ipn_callback():
    """Handle IPN (Instant Payment Notification) callbacks from NOWPayments with signature verification"""
    if not nowpayments:
        return jsonify({"error": "Crypto payments not configured"}), 503
    
    if not DB_AVAILABLE:
        logger.error("Database not available for IPN processing")
        # Return 503 to trigger NOWPayments retry mechanism
        return 'Database unavailable', 503
    
    try:
        received_signature = request.headers.get('x-nowpayments-sig')
        
        if not received_signature:
            logger.error("IPN callback missing x-nowpayments-sig header")
            return 'Missing signature', 401
        
        if not NOWPAYMENTS_IPN_SECRET:
            logger.error("NOWPAYMENTS_IPN_SECRET not configured - cannot verify IPN")
            return 'Configuration error', 500
        
        # Get raw request body as bytes for signature verification
        # CRITICAL: Use cache=True to preserve body for JSON parsing
        raw_body_bytes = request.get_data(cache=True)
        
        if not verify_nowpayments_ipn(NOWPAYMENTS_IPN_SECRET, raw_body_bytes, received_signature):
            logger.error("IPN signature verification failed - rejecting callback")
            return 'Invalid signature', 401
        
        logger.info("IPN signature verified successfully")
        
        # Parse JSON from the same bytes after verification
        data = json.loads(raw_body_bytes.decode('utf-8'))
        logger.info(f"Received crypto IPN callback: {data}")
        
        payment_id = data.get('payment_id')
        payment_status = data.get('payment_status')
        order_id = data.get('order_id')
        
        if not payment_id:
            logger.error("IPN callback missing payment_id")
            return 'Error', 400
        
        # Lock the payment record for atomic update (prevents race conditions)
        crypto_payment = CryptoPayment.query.filter_by(payment_id=payment_id).with_for_update().first()
        
        if not crypto_payment:
            logger.error(f"Crypto payment {payment_id} not found in database - requesting retry")
            # Return 503 to trigger NOWPayments retry (payment might be created soon)
            return 'Payment not found', 503
        
        if order_id and crypto_payment.order_id != order_id:
            logger.error(f"Order ID mismatch: expected {crypto_payment.order_id}, got {order_id}")
            return 'Order ID mismatch', 400
        
        # Validate price amount matches expected (in USD/fiat)
        if 'price_amount' in data:
            reported_price = float(data.get('price_amount'))
            if abs(reported_price - crypto_payment.price_amount) > 0.01:
                logger.error(f"Price amount mismatch: expected ${crypto_payment.price_amount}, got ${reported_price}")
                return 'Price mismatch', 400
        
        # Validate pay currency matches
        if 'pay_currency' in data:
            reported_currency = data.get('pay_currency', '').upper()
            if crypto_payment.pay_currency and reported_currency != crypto_payment.pay_currency:
                logger.error(f"Pay currency mismatch: expected {crypto_payment.pay_currency}, got {reported_currency}")
                return 'Currency mismatch', 400
        
        old_status = crypto_payment.payment_status
        crypto_payment.payment_status = payment_status
        
        logger.info(f"Crypto payment {payment_id} status updated from {old_status} to {payment_status}")
        
        # Handle partial/underpaid status explicitly
        if payment_status in ['partially_paid', 'underpaid']:
            logger.warning(f"Payment {payment_id} is {payment_status}, not crediting")
            db.session.commit()
            return 'OK', 200
        
        # Add credits only when payment is finished AND credits haven't been added yet (atomic idempotency)
        if payment_status == 'finished' and not crypto_payment.credits_added:
            user = User.query.get(crypto_payment.user_id)
            if user:
                # Validate actually paid amount (if provided)
                # Use outcome_amount (fiat) for comparison with price_amount
                if 'outcome_amount' in data:
                    actually_paid_fiat = float(data.get('outcome_amount'))
                    if actually_paid_fiat < (crypto_payment.price_amount - 0.01):
                        logger.error(f"Underpayment detected: expected ${crypto_payment.price_amount}, received ${actually_paid_fiat}")
                        db.session.commit()
                        return 'Underpayment', 400
                
                # Add credits to user account
                user.credits += crypto_payment.credits_purchased
                
                # Update last purchase timestamp to unlock video generation
                user.last_purchase_at = datetime.utcnow()
                
                # Create transaction record
                transaction = Transaction(
                    user_id=user.id,
                    credits_used=-crypto_payment.credits_purchased,
                    transaction_type='crypto_purchase',
                    description=f'Purchased {crypto_payment.credits_purchased} credits via {crypto_payment.pay_currency}'
                )
                db.session.add(transaction)
                
                # Mark as processed to prevent duplicate credits (idempotency)
                crypto_payment.credits_added = True
                crypto_payment.processed_at = datetime.utcnow()
                
                logger.info(f"Added {crypto_payment.credits_purchased} credits to user {user.telegram_id}. New balance: {user.credits}")
                
                # Commit before sending notification
                db.session.commit()
                
                # Send confirmation message after successful commit
                try:
                    confirmation_msg = f"✅ Payment confirmed! {crypto_payment.credits_purchased} credits have been added to your account.\n\nNew balance: {user.credits} credits"
                    send_message(user.telegram_id, confirmation_msg)
                except Exception as msg_error:
                    logger.error(f"Error sending confirmation message: {str(msg_error)}")
                
                return 'OK', 200
        elif payment_status == 'finished' and crypto_payment.credits_added:
            # Idempotency: Credits already added, just log and return success
            logger.info(f"Payment {payment_id} already processed at {crypto_payment.processed_at}. Skipping duplicate credit addition.")
            # Still commit to persist any status changes
            db.session.commit()
            return 'OK', 200
        
        elif payment_status == 'failed':
            logger.warning(f"Crypto payment {payment_id} failed")
            db.session.commit()
            try:
                user = User.query.get(crypto_payment.user_id)
                if user:
                    send_message(user.telegram_id, "❌ Payment failed. Please try again or contact support.")
            except Exception as msg_error:
                logger.error(f"Error sending failure message: {str(msg_error)}")
            return 'OK', 200
        
        db.session.commit()
        return 'OK', 200
        
    except Exception as e:
        logger.error(f"Error processing crypto IPN callback: {str(e)}")
        db.session.rollback()
        # Return 503 to trigger NOWPayments retry mechanism
        return 'Server error - retry requested', 503

@app.route('/api/crypto/payment-status/<payment_id>', methods=['GET'])
def get_crypto_payment_status(payment_id):
    """Check status of a crypto payment"""
    if not nowpayments:
        return jsonify({"error": "Crypto payments not configured"}), 503
    
    if not DB_AVAILABLE:
        return jsonify({"error": "Database not available"}), 503
    
    try:
        # Get from database
        crypto_payment = CryptoPayment.query.filter_by(payment_id=payment_id).first()
        
        if not crypto_payment:
            return jsonify({"error": "Payment not found"}), 404
        
        # Also check with NOWPayments API for latest status
        try:
            api_status = nowpayments.payment_status(payment_id)
            
            # Update database if status changed
            if api_status.get('payment_status') != crypto_payment.payment_status:
                crypto_payment.payment_status = api_status.get('payment_status')
                db.session.commit()
        except Exception as api_error:
            logger.error(f"Error fetching status from NOWPayments API: {str(api_error)}")
        
        return jsonify({
            "success": True,
            "payment_id": crypto_payment.payment_id,
            "payment_status": crypto_payment.payment_status,
            "pay_address": crypto_payment.pay_address,
            "pay_amount": crypto_payment.pay_amount,
            "pay_currency": crypto_payment.pay_currency,
            "credits_purchased": crypto_payment.credits_purchased
        }), 200
        
    except Exception as e:
        logger.error(f"Error checking crypto payment status: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/payment-history')
def payment_history():
    """View payment history for a user (debugging endpoint)"""
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Database not available",
            "message": "Payment history requires database connection"
        }), 503
    
    telegram_id = request.args.get('telegram_id')
    
    if not telegram_id:
        return jsonify({"error": "telegram_id query parameter is required"}), 400
    
    try:
        telegram_id = int(telegram_id)
        user = User.query.filter_by(telegram_id=telegram_id).first()
        
        if not user:
            return jsonify({
                "error": "User not found",
                "telegram_id": telegram_id
            }), 404
        
        payments = Payment.query.filter_by(user_id=user.id).order_by(Payment.created_at.desc()).all()
        
        payment_list = []
        for payment in payments:
            payment_list.append({
                "id": payment.id,
                "amount_cents": payment.amount,
                "amount_dollars": payment.amount / 100,
                "credits_purchased": payment.credits_purchased,
                "status": payment.status,
                "created_at": payment.created_at.strftime("%Y-%m-%d %H:%M:%S") if payment.created_at else None,
                "completed_at": payment.completed_at.strftime("%Y-%m-%d %H:%M:%S") if payment.completed_at else None,
                "stripe_session_id": payment.stripe_session_id
            })
        
        return jsonify({
            "user": {
                "telegram_id": user.telegram_id,
                "username": user.username,
                "current_credits": user.credits
            },
            "payments": payment_list,
            "total_payments": len(payment_list)
        }), 200
        
    except ValueError:
        return jsonify({"error": "Invalid telegram_id format"}), 400
    except Exception as e:
        logger.error(f"Error retrieving payment history: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/admin/broadcast', methods=['POST'])
def broadcast_message():
    """Send a broadcast message to all users (excluding recent Stars purchasers)
    
    POST /admin/broadcast
    {
        "token": "ADMIN_EXPORT_TOKEN",
        "message": "Your announcement message",
        "exclude_stars_purchasers_today": true
    }
    """
    if not DB_AVAILABLE:
        return jsonify({
            "error": "Database not available",
            "message": "Broadcast requires database connection"
        }), 503
    
    # Verify admin token
    admin_token = request.json.get('token')
    expected_token = os.environ.get('ADMIN_EXPORT_TOKEN')
    
    if not expected_token or admin_token != expected_token:
        return jsonify({"error": "Unauthorized"}), 401
    
    message = request.json.get('message')
    exclude_stars_today = request.json.get('exclude_stars_purchasers_today', False)
    
    if not message:
        return jsonify({"error": "message field is required"}), 400
    
    try:
        # Get all users with telegram_id
        all_users = User.query.filter(User.telegram_id.isnot(None)).all()
        
        # Get users who purchased Stars today (if exclusion enabled)
        excluded_user_ids = set()
        if exclude_stars_today:
            from sqlalchemy import func
            today_purchasers = db.session.query(TelegramPayment.user_id).filter(
                func.date(TelegramPayment.created_at) == datetime.utcnow().date()
            ).distinct().all()
            excluded_user_ids = {user_id for (user_id,) in today_purchasers}
        
        # Filter eligible users
        eligible_users = [u for u in all_users if u.id not in excluded_user_ids]
        
        # Send messages with rate limiting (30 messages/second max)
        success_count = 0
        error_count = 0
        errors = []
        
        for i, user in enumerate(eligible_users):
            try:
                result = send_message(user.telegram_id, message)
                if result and result.get('ok'):
                    success_count += 1
                else:
                    error_count += 1
                    errors.append(f"User {user.telegram_id}: {result.get('description', 'Unknown error')}")
                
                # Rate limiting: 30 messages/second = ~33ms between messages
                if (i + 1) % 30 == 0:
                    time.sleep(1)
                    
            except Exception as e:
                error_count += 1
                errors.append(f"User {user.telegram_id}: {str(e)}")
        
        return jsonify({
            "success": True,
            "total_users": len(all_users),
            "excluded_users": len(excluded_user_ids),
            "eligible_users": len(eligible_users),
            "messages_sent": success_count,
            "errors": error_count,
            "error_details": errors[:10] if errors else []
        }), 200
        
    except Exception as e:
        logger.error(f"Error broadcasting message: {str(e)}")
        return jsonify({"error": str(e)}), 500

# ============================================================================
# ERROR HANDLERS - Catch all uncaught exceptions
# ============================================================================

@app.errorhandler(Exception)
def handle_exception(e):
    """Global exception handler - catches all uncaught exceptions while preserving HTTP semantics"""
    import traceback
    from werkzeug.exceptions import HTTPException
    
    # Re-raise HTTPExceptions to preserve proper status codes (401, 404, etc.)
    if isinstance(e, HTTPException):
        return e
    
    # Only log and handle actual unexpected exceptions
    logger.error(f"⚠️ UNCAUGHT EXCEPTION: {str(e)}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    
    # Return 500 for unexpected exceptions
    return jsonify({
        "error": "Internal server error",
        "message": str(e)
    }), 500

@app.errorhandler(404)
def not_found(e):
    """Handle 404 errors"""
    return jsonify({"error": "Not found"}), 404

@app.errorhandler(500)
def internal_error(e):
    """Handle 500 errors"""
    logger.error(f"Internal server error: {str(e)}")
    return jsonify({"error": "Internal server error"}), 500

if __name__ == '__main__':
    # Run the Flask application
    # Note: keepalive thread is already started in global scope
    app.run(host='0.0.0.0', port=5000, debug=True)
